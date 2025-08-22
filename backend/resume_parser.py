import os
import logging
import tempfile
from pathlib import Path
import docx

# ✅ Correct import for pdfminer.six
from pdfminer.high_level import extract_text

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def parse_resume(file_or_path) -> str:
    """
    Parse resume from PDF, DOCX or TXT.

    Accepts either:
    - a filesystem path (str / os.PathLike)
    - a Streamlit UploadedFile (or any file-like with .read / .getbuffer and .name)
    """
    try:
        # 1) If it's an UploadedFile / file-like -> write to temp path first
        if hasattr(file_or_path, "read") and hasattr(file_or_path, "name"):
            suffix = Path(file_or_path.name).suffix.lower() or ".pdf"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                try:
                    buf = file_or_path.getbuffer() if hasattr(file_or_path, "getbuffer") else file_or_path.read()
                    tmp.write(buf if isinstance(buf, (bytes, bytearray)) else bytes(buf))
                except Exception:
                    # Fallback: read() then write
                    tmp.write(file_or_path.read())
                tmp_path = tmp.name

            try:
                text = _parse_resume_path(tmp_path)
                return text
            finally:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

        # 2) Else treat as a path
        elif isinstance(file_or_path, (str, os.PathLike)):
            return _parse_resume_path(str(file_or_path))

        else:
            raise ValueError("Unsupported input: provide a file path or an UploadedFile object.")

    except Exception as e:
        logger.error(f"❌ Error parsing resume: {e}")
        raise ValueError(f"Failed to parse resume: {str(e)}")

def _parse_resume_path(file_path: str) -> str:
    """Internal: parse by path (original behavior)."""
    if not os.path.exists(file_path):
        raise ValueError(f"File not found: {file_path}")

    file_extension = Path(file_path).suffix.lower()
    logger.info(f"📄 Parsing resume: {os.path.basename(file_path)} ({file_extension})")

    if file_extension == '.pdf':
        return parse_pdf(file_path)
    elif file_extension == '.docx':
        return parse_docx(file_path)
    elif file_extension == '.txt':
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: PDF, DOCX, TXT")

def parse_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfminer.six."""
    try:
        text = extract_text(file_path)
        if not text or not text.strip():
            logger.warning("⚠️ No text extracted. PDF might be image-based or corrupted.")
            raise ValueError("PDF appears to be image-based or corrupted. Please provide a text-based PDF.")
        return text
    except Exception as e:
        raise ValueError(f"PDF parsing failed: {str(e)}")


def parse_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        text = "\n".join(paragraphs + tables_text)
        if not text.strip():
            raise ValueError("DOCX file appears to be empty or contains no readable text.")
        return text
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {str(e)}")

def parse_txt(file_path: str) -> str:
    """Extract text from TXT."""
    try:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
        if not text.strip():
            raise ValueError("TXT file appears to be empty.")
        return text
    except Exception as e:
        raise ValueError(f"TXT parsing failed: {str(e)}")

def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    lines = []
    for line in text.split('\n'):
        cleaned_line = ' '.join(line.split())
        if cleaned_line:
            lines.append(cleaned_line)
    import re
    cleaned_text = '\n'.join(lines)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    return cleaned_text.strip()

def save_parsed_text(original_file_path: str, text: str) -> None:
    """Save parsed text for reference."""
    try:
        output_dir = "outputs/parsed_resumes"
        os.makedirs(output_dir, exist_ok=True)
        original_name = Path(original_file_path).stem
        output_filename = f"{original_name}_parsed.txt"
        output_path = os.path.join(output_dir, output_filename)
        from datetime import datetime as _dt
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("=== PARSED RESUME TEXT ===\n\n")
            f.write(f"Original file: {os.path.basename(original_file_path)}\n")
            f.write(f"Parsed on: {_dt.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Text length: {len(text)} characters\n")
            f.write("\n" + "="*50 + "\n\n")
            f.write(text)
        logger.info(f"💾 Parsed text saved to: {output_path}")
    except Exception as e:
        logger.warning(f"⚠️ Could not save parsed text: {e}")

# resume_parser.py
import io
from typing import Optional
from PyPDF2 import PdfReader

def extract_text_from_pdf(file_bytes: bytes) -> Optional[str]:
    """
    Extract text from a PDF file given as bytes.
    Returns the extracted text or None if parsing fails.
    """
    try:
        text = ""
        with io.BytesIO(file_bytes) as pdf_file:
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        print(f"[ERROR] Failed to extract text from PDF: {e}")
        return None
